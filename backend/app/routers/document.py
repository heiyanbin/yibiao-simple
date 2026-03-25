"""文档处理相关API路由"""
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Query, Request
from fastapi.responses import StreamingResponse
from ..models.schemas import FileUploadResponse, AnalysisRequest, AnalysisType, WordExportRequest
from ..services.file_service import FileService
from ..services.openai_service import OpenAIService
from ..utils.system_config import system_config
from ..utils.prompt_manager import get_default_overview_prompt, get_default_requirements_prompt
from ..utils.sse import sse_response
from ..routers.auth import get_current_user_from_request
from ..models.db_models import Job
from ..utils.database import get_session
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
import json
import io
import os
import re
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from urllib.parse import quote

router = APIRouter(prefix="/api/document", tags=["文档处理"])


def set_run_font_simsun(run: docx.text.run.Run) -> None:
    """统一将 run 字体设置为宋体（包含 EastAsia 字体设置）"""
    run.font.name = "宋体"
    r = run._element.rPr
    if r is not None and r.rFonts is not None:
        r.rFonts.set(qn("w:eastAsia"), "宋体")


def set_paragraph_font_simsun(paragraph: docx.text.paragraph.Paragraph) -> None:
    """将段落内所有 runs 字体设置为宋体"""
    for run in paragraph.runs:
        set_run_font_simsun(run)


@router.post("/upload", response_model=FileUploadResponse)
async def upload_file(
    file: UploadFile = File(...),
    job_id: int = Query(None),
    session: AsyncSession = Depends(get_session),
    current_user = Depends(get_current_user_from_request)
):
    """上传文档文件并提取文本内容（需要登录）"""
    try:
        # 检查文件类型
        allowed_types = [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]

        if file.content_type not in allowed_types:
            return FileUploadResponse(
                success=False,
                message="不支持的文件类型，请上传PDF或Word文档"
            )

        # 处理文件并提取文本，返回文件路径和文本内容
        file_path, file_content = await FileService.process_uploaded_file(file)

        # 如果提供了 job_id，更新任务记录
        if job_id:
            result = await session.execute(
                select(Job).where(Job.id == job_id, Job.user_id == current_user.id)
            )
            job = result.scalar_one_or_none()
            if job:
                job.source_file_name = file.filename
                job.source_file_path = file_path  # 保存文件路径
                job.file_content = file_content
                await session.commit()

        return FileUploadResponse(
            success=True,
            message=f"文件 {file.filename} 上传成功",
            file_content=file_content
        )

    except Exception as e:
        return FileUploadResponse(
            success=False,
            message=f"文件处理失败: {str(e)}"
        )


@router.get("/download/{job_id}")
async def download_source_file(
    job_id: int,
    request: Request,
    session: AsyncSession = Depends(get_session)
):
    """下载任务的原始文件"""
    user = await get_current_user_from_request(request)

    result = await session.execute(
        select(Job).where(Job.id == job_id, Job.user_id == user.id)
    )
    job = result.scalar_one_or_none()

    if not job:
        raise HTTPException(status_code=404, detail="任务不存在")

    if not job.source_file_path or not os.path.exists(job.source_file_path):
        raise HTTPException(status_code=404, detail="文件不存在或已被删除")

    from fastapi.responses import FileResponse
    return FileResponse(
        job.source_file_path,
        filename=job.source_file_name,
        media_type="application/octet-stream"
    )


@router.post("/analyze-stream")
async def analyze_document_stream(
    request: AnalysisRequest,
    current_user = Depends(get_current_user_from_request)
):
    """流式分析文档内容（需要登录）"""
    try:
        # 检查系统是否已配置 API Key
        if not system_config.is_configured():
            raise HTTPException(status_code=400, detail="系统未配置 API Key，请联系管理员")

        # 创建OpenAI服务实例，使用请求中的模型或默认模型
        model_name = getattr(request, 'model_name', None)
        openai_service = OpenAIService(model_name=model_name)
        
        async def generate():
            # 构建分析提示词
            if request.custom_prompt:
                # 使用自定义提示词
                system_prompt = request.custom_prompt
            elif request.analysis_type == AnalysisType.OVERVIEW:
                system_prompt = get_default_overview_prompt()
            else:  # requirements
                system_prompt = get_default_requirements_prompt()

            analysis_type_cn = "项目概述" if request.analysis_type == AnalysisType.OVERVIEW else "技术评分要求"
            user_prompt = f"请分析以下招标文件内容，提取{analysis_type_cn}信息：\n\n{request.file_content}"
            
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]
            
            # 流式返回分析结果
            async for chunk in openai_service.stream_chat_completion(messages, temperature=0.3):
                yield f"data: {json.dumps({'chunk': chunk}, ensure_ascii=False)}\n\n"
            
            # 发送结束信号
            yield "data: [DONE]\n\n"
        
        return sse_response(generate())
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文档分析失败: {str(e)}")


@router.post("/export-word")
async def export_word(
    request: WordExportRequest,
    current_user = Depends(get_current_user_from_request)
):
    """根据目录数据导出Word文档（需要登录）"""
    try:
        doc = docx.Document()

        # 统一设置文档的基础字体为宋体，取消普通段落默认加粗
        try:
            styles = doc.styles
            base_styles = ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title"]
            for style_name in base_styles:
                if style_name in styles:
                    style = styles[style_name]
                    font = style.font
                    font.name = "宋体"
                    # 设置中文字体
                    if style._element.rPr is None:
                        style._element._add_rPr()
                    rpr = style._element.rPr
                    rpr.rFonts.set(qn("w:eastAsia"), "宋体")
                    if style_name == "Normal":
                        font.bold = False
        except Exception:
            # 字体设置失败不影响文档生成，忽略
            pass

        # AI 生成声明
        p = doc.add_paragraph()
        run = p.add_run("内容由AI生成")
        run.italic = True
        run.font.size = Pt(9)
        set_run_font_simsun(run)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 文档标题
        title = request.project_name or "投标技术文件"
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        set_run_font_simsun(title_run)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 项目概述
        if request.project_overview:
            heading = doc.add_heading("项目概述", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_font_simsun(heading)
            overview_p = doc.add_paragraph(request.project_overview)
            set_paragraph_font_simsun(overview_p)
            overview_p_format = overview_p.paragraph_format
            overview_p_format.space_after = Pt(12)

        # 简单的 Markdown 段落解析：支持标题、列表、表格和基础加粗/斜体
        def add_markdown_runs(para: docx.text.paragraph.Paragraph, text: str) -> None:
            """在指定段落中追加 markdown 文本的 runs"""
            pattern = r"(\*\*.*?\*\*|\*.*?\*|`.*?`)"
            parts = re.split(pattern, text)
            for part in parts:
                if not part:
                    continue
                run = para.add_run()
                # 加粗
                if part.startswith("**") and part.endswith("**") and len(part) > 4:
                    run.text = part[2:-2]
                    run.bold = True
                # 斜体
                elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                    run.text = part[1:-1]
                    run.italic = True
                # 行内代码：这里只去掉反引号
                elif part.startswith("`") and part.endswith("`") and len(part) > 2:
                    run.text = part[1:-1]
                else:
                    run.text = part
                # 确保字体为宋体
                set_run_font_simsun(run)

        def add_markdown_paragraph(text: str) -> None:
            """将一段 Markdown 文本解析为一个普通段落，保留加粗/斜体效果"""
            para = doc.add_paragraph()
            add_markdown_runs(para, text)
            para.paragraph_format.space_after = Pt(6)

        def parse_markdown_blocks(content: str):
            """
            识别 Markdown 内容中的块级元素，返回结构化的 block 列表：
            - ('list', items)        items: [(kind, num_str, text), ...]
            - ('table', rows)        rows: [text, ...]
            - ('heading', level, text)
            - ('paragraph', text)
            """
            blocks = []
            lines = content.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i].rstrip("\r").strip()
                if not line:
                    i += 1
                    continue

                # 列表项（有序/无序）
                if line.startswith("- ") or line.startswith("* ") or re.match(r"^\d+\.\s", line):
                    # items: (kind, number, text)
                    items = []
                    while i < len(lines):
                        raw = lines[i].rstrip("\r")
                        stripped = raw.strip()
                        # 无序列表
                        if stripped.startswith("- ") or stripped.startswith("* "):
                            text = re.sub(r"^[-*]\s+", "", stripped).strip()
                            if text:
                                items.append(("unordered", None, text))
                            i += 1
                            continue
                        # 有序列表（1. xxx）
                        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
                        if m_num:
                            num_str, text = m_num.groups()
                            text = text.strip()
                            if text:
                                items.append(("ordered", num_str, text))
                            i += 1
                            continue
                        break

                    if items:
                        blocks.append(("list", items))
                    continue

                # 表格（简化为每行一个段落，单元格用 | 分隔）
                if "|" in line:
                    rows = []
                    while i < len(lines):
                        raw = lines[i].rstrip("\r")
                        stripped = raw.strip()
                        if "|" in stripped:
                            # 跳过仅由 - 和 | 组成的分隔行
                            if not re.match(r"^\|?[-\s\|]+\|?$", stripped):
                                cells = [c.strip() for c in stripped.split("|")]
                                row_text = " | ".join([c for c in cells if c])
                                if row_text:
                                    rows.append(row_text)
                            i += 1
                        else:
                            break
                    if rows:
                        blocks.append(("table", rows))
                    continue

                # Markdown 标题（# / ## / ###）
                if line.startswith("#"):
                    m = re.match(r"^(#+)\s*(.*)$", line)
                    if m:
                        level_marks, title_text = m.groups()
                        level = min(len(level_marks), 3)
                        blocks.append(("heading", level, title_text.strip()))
                    i += 1
                    continue

                # 普通段落：合并连续的普通行
                para_lines = []
                while i < len(lines):
                    raw = lines[i].rstrip("\r")
                    stripped = raw.strip()
                    if (
                        stripped
                        and not stripped.startswith("-")
                        and not stripped.startswith("*")
                        and "|" not in stripped
                        and not stripped.startswith("#")
                    ):
                        para_lines.append(stripped)
                        i += 1
                    else:
                        break
                if para_lines:
                    text = " ".join(para_lines)
                    blocks.append(("paragraph", text))
                else:
                    i += 1

            return blocks

        def render_markdown_blocks(blocks) -> None:
            """将结构化的 Markdown blocks 渲染到文档"""
            for block in blocks:
                kind = block[0]
                if kind == "list":
                    items = block[1]
                    for item_kind, num_str, text in items:
                        p = doc.add_paragraph()
                        if item_kind == "unordered":
                            # 使用"• "模拟项目符号
                            run = p.add_run("• ")
                            set_run_font_simsun(run)
                        else:
                            # 有序列表：输出 "1. " 这样的前缀
                            prefix = f"{num_str}."
                            run = p.add_run(prefix + " ")
                            set_run_font_simsun(run)
                        # 紧跟在同一段落中追加列表文本
                        add_markdown_runs(p, text)
                elif kind == "table":
                    rows = block[1]
                    for row in rows:
                        add_markdown_paragraph(row)
                elif kind == "heading":
                    _, level, text = block
                    heading = doc.add_heading(text, level=level)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    set_paragraph_font_simsun(heading)
                elif kind == "paragraph":
                    _, text = block
                    add_markdown_paragraph(text)

        def add_markdown_content(content: str) -> None:
            """解析并渲染 Markdown 文本到文档"""
            blocks = parse_markdown_blocks(content)
            render_markdown_blocks(blocks)

        # 递归构建文档内容（章节和内容）
        def add_outline_items(items, level: int = 1):
            for item in items:
                # 章节标题
                if level <= 3:
                    heading = doc.add_heading(f"{item.id} {item.title}", level=level)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for hr in heading.runs:
                        hr.font.name = "宋体"
                        rr = hr._element.rPr
                        if rr is not None and rr.rFonts is not None:
                            rr.rFonts.set(qn("w:eastAsia"), "宋体")
                else:
                    para = doc.add_paragraph()
                    run = para.add_run(f"{item.id} {item.title}")
                    run.bold = True
                    run.font.name = "宋体"
                    rr = run._element.rPr
                    if rr is not None and rr.rFonts is not None:
                        rr.rFonts.set(qn("w:eastAsia"), "宋体")
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(3)

                # 叶子节点内容
                if not item.children:
                    content = item.content or ""
                    if content.strip():
                        add_markdown_content(content)
                else:
                    add_outline_items(item.children, level + 1)

        add_outline_items(request.outline)

        # 输出到内存并返回
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{request.project_name or '标书文档'}.docx"
        # 使用 RFC 5987 格式对文件名进行 URL 编码，避免非 ASCII 字符导致的编码错误
        encoded_filename = quote(filename)
        content_disposition = f"attachment; filename*=UTF-8''{encoded_filename}"
        headers = {
            "Content-Disposition": content_disposition
        }

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
    except Exception as e:
        # 打印详细错误信息到控制台，方便排查
        import traceback
        print("导出Word失败:", str(e))
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"导出Word失败: {str(e)}")